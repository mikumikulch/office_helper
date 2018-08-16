#!/usr/bin/python3
# -*- coding:utf-8 -*-

"""
自动填写加班单、并保存到对应文件目录中。
最后发送回邮件到邮箱、提示打印。打印功能也许可以用 applescript+pages 实现？
~~~~~~~~~~~~~~~~~~~~~~~~~~~~
"""
import os

__author__ = 'Chuck Lin'

from datetime import datetime

import docx
import logging

from attendance_spider import HelperConfig

from docx.oxml.ns import qn
from docx.shared import Pt

from mail_sender import email_sender

logger_name = 'office_helper'
logger = logging.getLogger(logger_name)

__author__ = 'Chuck Lin'

g_doc = docx.Document('make_document/document/加班审批表.docx')
style = g_doc.styles['Normal']
style.font.name = u'宋体'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


def ot_file_remover(x):
    if len(x) > 13:
        os.remove('./document/%s' % x)


# 备用符号 ☑
def reset_form_date_paragraphs(flg):
    if flg:
        # 加班审批表的第一个填表日期的段落数
        paragraph_number = 1
    else:
        # 加班审批表的第二个填表日期的段落数
        paragraph_number = 5
    # 清楚段落的内容
    g_doc.paragraphs[paragraph_number].clear()
    # 获取填表日期段落
    form_date_paragraphs = g_doc.paragraphs[paragraph_number]
    # form_date_paragraphs.clear()
    now = datetime.now()

    # 填写填表日期
    mutipart_date = '填表日期%d年%d月%d日' % (now.year, now.month, now.day)
    run = form_date_paragraphs.add_run(mutipart_date, 'Default Paragraph Font')
    run.font.size = Pt(10)
    return g_doc


def write_info_to_memory(checkout_date, checkout_time):
    logger.info('开始自动填写加班模板表。checkout_date -> %s checkout_time -> %s' % (checkout_date, checkout_time))

    # 获取加班单中的表格集合
    def set_tables_fontsize(x):
        x.style.font.size = Pt(11)
        return x

    pre_tables = [table for table in g_doc.tables]
    tables = list(map(set_tables_fontsize, pre_tables))

    apply_name = tables[0].rows[0].cells[1].text
    apply_name_2 = tables[1].rows[0].cells[1].text
    if apply_name_2 != '　':
        logger.info('所有表单已经被使用，建立新的加班审批表填写内容')
        flg = True
        local_doc = docx.Document('make_document/document/加班审批表_模板.docx')
        pre_tables = [table for table in local_doc.tables]
        tables = list(map(set_tables_fontsize, pre_tables))
        table = tables[0]
    elif apply_name != '　':
        logger.info('上部分表单已经被使用，使用下半部分表单的内容')
        flg = False
        table = tables[1]
    else:
        logger.info('未被使用的表单。开始填入数据')
        flg = True
        table = tables[0]
    # 填写加班日期填写时间项
    reset_form_date_paragraphs(flg)  # 填写加班事由
    # TODO 自动爬取近期的需求，
    table.rows[1].cells[2].text = '真量贷中新需求开发与系统遗留 bug 修正'
    # 填写申请人
    table.rows[0].cells[1].text = HelperConfig.user_name
    # 填写所在部门
    table.rows[0].cells[3].text = '信息技术部门'
    table.rows[0].cells[5].text = '程序员'
    # 填写加班类型
    # TODO 根据加班时间的不同，填入不同的加班类型
    table.rows[4].cells[2].text = '☑ 延时加班    □ 休息日加班    □ 法定节假日加班'
    # 申请加班时间
    year, month, day, hour = checkout_date.year, checkout_date.month, checkout_date.day, checkout_time.hour
    apply_work_time = '%d 年 %d 月 %d 日  19 时 至  %d 年 %d  月  %d 日  %d 时' % (year, month, day, year, month, day, hour)
    table.rows[5].cells[2].text = apply_work_time
    return flg


def write_document(checkout_date, checkout_time):
    flg = write_info_to_memory(checkout_date, checkout_time)
    save_file_path = 'make_document/document/加班审批表.docx'
    create_file_path = 'make_document/document/加班审批表_%s.docx' % (datetime.now().strftime('%Y-%m-%d'))
    # 为防止留存的加班表过度，当文档数量大于30时，删除除模板文档以外的所有文档。
    docx_list = [x for x in os.listdir('make_document/document') if os.path.splitext(x)[1] == '.docx']
    if len(docx_list) > 30:
        logger.info('当前留存的加班审批表过多。删除模板以外的所有文档')
        list(map(ot_file_remover, docx_list))

    # 由于加班不一定是连续性的，所以每次加班都发送邮件并且覆盖加班审批表
    g_doc.save(create_file_path)
    attachment_name = '加班审批表_%s.docx' % datetime.now().strftime('%Y-%m-%d')
    email_sender.send_mail('【敏感词办公小助手】加班审批表_%s' % datetime.now().strftime('%Y-%m-%d'), create_file_path,
                           attachment_name)
    template_doc = docx.Document('make_document/document/加班审批表_模板.docx')
    template_doc.save(save_file_path)
    logger.info('自动写入加班审批表并且发送邮件完毕')

    # # 为 true 时在旧的文档上写入。
    # if flg:
    #     g_doc.save(save_file_path)
    #     logger.info('自动写入加班审批表处理结束')
    # # flg 为 false 时，代表目前文档已经写满。将会创建新的文档，并附上时间。并读取模板的内容覆盖掉目前的加班审批表
    #
    # else:
    #     g_doc.save(create_file_path)
    #     attachment_name = '加班审批表_%s.docx' % datetime.now().strftime('%Y-%m-%d')
    #     email_sender.send_mail('【敏感词办公小助手】加班审批表_%s' % datetime.now().strftime('%Y-%m-%d'), create_file_path,
    #                            attachment_name)
    #     template_doc = docx.Document('make_document/document/加班审批表_模板.docx')
    #     template_doc.save(save_file_path)
    #     logger.info('自动写入加班审批表并且发送邮件完毕')
